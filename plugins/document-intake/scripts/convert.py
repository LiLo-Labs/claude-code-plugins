#!/usr/bin/env python3
"""Stage 1 — Convert mixed-format documents to markdown.

Supports: PDF, DOCX, XLSX, CSV, PPTX, EML, VSDX. Extracts embedded images to a sibling
`embedded_images/<safe-name>/` folder and writes inline placeholders:

    **Embedded image (page N, image M):** `embedded_images/<safe>/pageN_imgM.png`

Usage:
    python3 convert.py --input <source-folder>
        [--output <processed-folder>]    # default: <source-folder>/processed

All later pipeline stages operate on the same `processed/` folder in place.
Final layout after all stages:

    <source-folder>/
        processed/
            <file>.md                    ← polished, chunk-marked
            ...
            embedded_images/
            _chunks/
            _polish_reports/
            _summary.md
            _summary.json
            _chunk_index.json
"""
from __future__ import annotations
import argparse, hashlib, io, json, pathlib, re, sys, zipfile
from email import policy
from email.parser import BytesParser

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
try:
    import docx as python_docx
except ImportError:
    python_docx = None
try:
    import openpyxl
except ImportError:
    openpyxl = None
try:
    from PIL import Image
except ImportError:
    Image = None


SAFE_RX = re.compile(r"[^A-Za-z0-9._-]+")

def safe_name(stem: str) -> str:
    return SAFE_RX.sub("_", stem).strip("_")


# ─── PDF ─────────────────────────────────────────────────────────────────────
def convert_pdf(src: pathlib.Path, img_dir: pathlib.Path) -> tuple[str, int]:
    if fitz is None:
        return f"[PDF conversion unavailable — install pymupdf]\n", 0
    doc = fitz.open(src)
    out = [f"# {src.name}", "", f"**Source file:** `{src.name}`",
           f"**Pages:** {doc.page_count}", "", "---", ""]
    img_count = 0
    seen_hashes: dict[str, str] = {}
    for i, page in enumerate(doc, start=1):
        out.append(f"## Page {i}")
        out.append("")
        out.append(page.get_text("text").rstrip())
        out.append("")
        for idx, info in enumerate(page.get_images(full=True), start=1):
            xref = info[0]
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.n > 4 or (pix.colorspace is not None and pix.colorspace.name not in ("DeviceRGB", "DeviceGray")):
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                data = pix.tobytes("png")
            except Exception:
                # Fall back to the raw embedded bytes.
                try:
                    info_dict = doc.extract_image(xref)
                    data = info_dict["image"]
                    ext = info_dict.get("ext", "png")
                except Exception as e:
                    out.append(f"*[image extraction failed (page {i}, image {idx}): {e}]*")
                    continue
                h = hashlib.sha1(data).hexdigest()
                if h in seen_hashes:
                    path = seen_hashes[h]
                else:
                    path = img_dir / f"page{i}_img{idx}.{ext}"
                    path.write_bytes(data)
                    seen_hashes[h] = str(path)
                    img_count += 1
                rel = pathlib.Path(path).relative_to(img_dir.parent.parent)  # include 'embedded_images/' prefix
                out.append(f"**Embedded image (page {i}, image {idx}):** `{rel}`")
                continue
            h = hashlib.sha1(data).hexdigest()
            if h in seen_hashes:
                path = seen_hashes[h]
            else:
                path = img_dir / f"page{i}_img{idx}.png"
                path.write_bytes(data)
                seen_hashes[h] = str(path)
                img_count += 1
            rel = pathlib.Path(path).relative_to(img_dir.parent.parent)  # include 'embedded_images/' prefix
            out.append(f"**Embedded image (page {i}, image {idx}):** `{rel}`")
        out.append("")
    doc.close()
    return "\n".join(out), img_count


# ─── DOCX ────────────────────────────────────────────────────────────────────
def convert_docx(src: pathlib.Path, img_dir: pathlib.Path) -> tuple[str, int]:
    if python_docx is None:
        return f"[DOCX conversion unavailable — install python-docx]\n", 0
    doc = python_docx.Document(src)
    out = [f"# {src.name}", "", f"**Source file:** `{src.name}`", "", "---", ""]
    # walk body elements in order; add headings by style name
    for p in doc.paragraphs:
        style = (p.style.name or "").lower()
        text = p.text.rstrip()
        if not text:
            out.append("")
            continue
        if style.startswith("heading 1"):
            out.append(f"## {text}")
        elif style.startswith("heading 2"):
            out.append(f"### {text}")
        elif style.startswith("heading 3"):
            out.append(f"#### {text}")
        else:
            out.append(text)
        out.append("")
    # tables
    for ti, table in enumerate(doc.tables, 1):
        out.append(f"### Table {ti}")
        out.append("")
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        out.append("| " + " | ".join(rows[0]) + " |")
        out.append("|" + "|".join(["---"] * len(rows[0])) + "|")
        for r in rows[1:]:
            out.append("| " + " | ".join(r) + " |")
        out.append("")
    # media
    img_count = 0
    with zipfile.ZipFile(src) as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                data = z.read(name)
                ext = pathlib.Path(name).suffix or ".bin"
                out_name = f"image{img_count+1}{ext}"
                (img_dir / out_name).write_bytes(data)
                rel = pathlib.Path(img_dir / out_name).relative_to(img_dir.parent.parent)  # include 'embedded_images/' prefix
                out.append(f"**Embedded image (image {img_count+1}):** `{rel}`")
                img_count += 1
        out.append("")
    return "\n".join(out), img_count


# ─── XLSX ────────────────────────────────────────────────────────────────────
def convert_xlsx(src: pathlib.Path, _img_dir: pathlib.Path) -> tuple[str, int]:
    if openpyxl is None:
        return f"[XLSX conversion unavailable — install openpyxl]\n", 0
    wb = openpyxl.load_workbook(src, data_only=True, read_only=True)
    out = [f"# {src.name}", "", f"**Source file:** `{src.name}`",
           f"**Sheets:** {len(wb.sheetnames)}", "", "---", ""]
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        out.append(f"## Sheet: {sheet_name}")
        out.append("")
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            out.append("_(empty)_"); out.append(""); continue
        # pick first non-empty row as header
        hdr = [str(c) if c is not None else "" for c in rows[0]]
        out.append("| " + " | ".join(hdr) + " |")
        out.append("|" + "|".join(["---"] * len(hdr)) + "|")
        for r in rows[1:]:
            vals = [str(c) if c is not None else "" for c in r]
            out.append("| " + " | ".join(vals) + " |")
        out.append("")
    return "\n".join(out), 0


# ─── CSV ─────────────────────────────────────────────────────────────────────
def convert_csv(src: pathlib.Path, _img_dir: pathlib.Path) -> tuple[str, int]:
    import csv
    out = [f"# {src.name}", "", f"**Source file:** `{src.name}`", "", "---", ""]
    # Auto-detect delimiter; default to comma if sniff fails.
    with src.open(newline="", encoding="utf-8-sig", errors="replace") as fh:
        sample = fh.read(4096); fh.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        reader = csv.reader(fh, dialect)
        rows = list(reader)
    if not rows:
        out.append("_(empty CSV)_"); return "\n".join(out), 0
    # Normalize row widths.
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    # Pipe-escape cell contents.
    esc = lambda s: s.replace("|", r"\|").replace("\n", " ").strip()
    out.append("| " + " | ".join(esc(c) for c in rows[0]) + " |")
    out.append("|" + "|".join(["---"] * width) + "|")
    for r in rows[1:]:
        out.append("| " + " | ".join(esc(c) for c in r) + " |")
    out.append("")
    return "\n".join(out), 0


# ─── EML ─────────────────────────────────────────────────────────────────────
def convert_eml(src: pathlib.Path, _img_dir: pathlib.Path) -> tuple[str, int]:
    msg = BytesParser(policy=policy.default).parse(src.open("rb"))
    out = [f"# {src.name}", "", f"**Source file:** `{src.name}`", ""]
    for h in ("From", "To", "Cc", "Subject", "Date"):
        v = msg.get(h)
        if v:
            out.append(f"- **{h}:** {v}")
    out.append("")
    out.append("---")
    out.append("")
    # prefer text/plain
    body = None
    if msg.is_multipart():
        for part in msg.walk():
            ctype = part.get_content_type()
            if ctype == "text/plain" and body is None:
                body = part.get_content()
        if body is None:
            for part in msg.walk():
                if part.get_content_type() == "text/html":
                    body = re.sub(r"<[^>]+>", "", part.get_content())
                    break
    else:
        body = msg.get_content()
    if body:
        out.append(body.strip())
    out.append("")
    return "\n".join(out), 0


# ─── PPTX ────────────────────────────────────────────────────────────────────
def convert_pptx(src: pathlib.Path, img_dir: pathlib.Path) -> tuple[str, int]:
    out = [f"# {src.name}", "", f"**Source file:** `{src.name}`", "", "---", ""]
    img_count = 0
    try:
        with zipfile.ZipFile(src) as z:
            names = z.namelist()
            slides = sorted(
                (n for n in names if n.startswith("ppt/slides/slide") and n.endswith(".xml")),
                key=lambda x: int(re.search(r"slide(\d+)\.xml$", x).group(1) if re.search(r"slide(\d+)\.xml$", x) else "0"),
            )
            notes_map = {}
            for n in names:
                m = re.match(r"ppt/notesSlides/notesSlide(\d+)\.xml$", n)
                if m:
                    notes_map[int(m.group(1))] = n

            def text_runs(xml: str) -> list[str]:
                # Each <a:t> holds one run of slide text; concatenate consecutive runs
                # in the same <a:p> into one paragraph for readability.
                paras = re.findall(r"<a:p\b[^>]*>(.*?)</a:p>", xml, flags=re.DOTALL)
                lines = []
                for p in paras:
                    runs = re.findall(r"<a:t[^>]*>(.*?)</a:t>", p, flags=re.DOTALL)
                    decoded = []
                    for r in runs:
                        # Strip any leftover inline tags, decode XML entities.
                        r = re.sub(r"<[^>]+>", "", r)
                        r = (r.replace("&amp;", "&").replace("&lt;", "<")
                              .replace("&gt;", ">").replace("&quot;", "\"")
                              .replace("&apos;", "'"))
                        decoded.append(r)
                    line = "".join(decoded).strip()
                    if line:
                        lines.append(line)
                return lines

            for slide_path in slides:
                m = re.search(r"slide(\d+)\.xml$", slide_path)
                slide_num = int(m.group(1)) if m else 0
                xml = z.read(slide_path).decode("utf-8", errors="replace")
                lines = text_runs(xml)
                out.append(f"## Slide {slide_num}")
                out.append("")
                if lines:
                    for ln in lines:
                        out.append(ln)
                else:
                    out.append("_(no text on this slide)_")
                out.append("")
                if slide_num in notes_map:
                    notes_xml = z.read(notes_map[slide_num]).decode("utf-8", errors="replace")
                    notes = text_runs(notes_xml)
                    # Drop the boilerplate slide number that PPTX puts in notes.
                    notes = [n for n in notes if not re.match(r"^\d+\s*$", n)]
                    if notes:
                        out.append("### Speaker notes")
                        out.append("")
                        for n in notes:
                            out.append(n)
                        out.append("")

            # Images from ppt/media/
            for name in sorted(names):
                if name.startswith("ppt/media/"):
                    data = z.read(name)
                    ext = pathlib.Path(name).suffix or ".bin"
                    img_count += 1
                    out_name = f"image{img_count}{ext}"
                    (img_dir / out_name).write_bytes(data)
                    rel = pathlib.Path(img_dir / out_name).relative_to(img_dir.parent.parent)
                    out.append(f"**Embedded image (image {img_count}):** `{rel}`")
            out.append("")
    except zipfile.BadZipFile:
        out.append("_(PPTX file unreadable — not a valid ZIP)_")
    return "\n".join(out), img_count


# ─── VSDX ────────────────────────────────────────────────────────────────────
def convert_vsdx(src: pathlib.Path, img_dir: pathlib.Path) -> tuple[str, int]:
    out = [f"# {src.name}", "", f"**Source file:** `{src.name}`", "", "---", ""]
    img_count = 0
    try:
        with zipfile.ZipFile(src) as z:
            page_xmls = sorted(n for n in z.namelist() if n.startswith("visio/pages/page") and n.endswith(".xml"))
            for pn, pxml in enumerate(page_xmls, start=1):
                xml = z.read(pxml).decode("utf-8", errors="replace")
                texts = re.findall(r"<Text[^>]*>(.*?)</Text>", xml, flags=re.DOTALL)
                texts = [re.sub(r"<[^>]+>", "", t).strip() for t in texts]
                out.append(f"## Page {pn}")
                out.append("")
                for t in texts:
                    if t: out.append(t)
                out.append("")
            # media
            for name in z.namelist():
                if name.startswith("visio/media/"):
                    data = z.read(name)
                    ext = pathlib.Path(name).suffix or ".bin"
                    out_name = f"image{img_count+1}{ext}"
                    (img_dir / out_name).write_bytes(data)
                    rel = pathlib.Path(img_dir / out_name).relative_to(img_dir.parent.parent)  # include 'embedded_images/' prefix
                    out.append(f"**Embedded image (image {img_count+1}):** `{rel}`")
                    img_count += 1
            out.append("")
    except zipfile.BadZipFile:
        out.append("_(VSDX file unreadable — not a valid ZIP)_")
    return "\n".join(out), img_count


# ─── DRIVER ──────────────────────────────────────────────────────────────────
CONVERTERS = {
    ".pdf":  convert_pdf,
    ".docx": convert_docx,
    ".xlsx": convert_xlsx,
    ".csv":  convert_csv,
    ".pptx": convert_pptx,
    ".eml":  convert_eml,
    ".vsdx": convert_vsdx,
}

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input",  required=True, help="Folder containing source docs")
    ap.add_argument("--output", default=None,  help="Output folder (default: <input>/convert)")
    args = ap.parse_args()

    src_root = pathlib.Path(args.input).expanduser().resolve()
    if not src_root.is_dir():
        print(f"Input not a directory: {src_root}", file=sys.stderr); sys.exit(2)
    out_root = pathlib.Path(args.output).expanduser().resolve() if args.output else src_root / "processed"
    out_root.mkdir(exist_ok=True)
    img_root = out_root / "embedded_images"
    img_root.mkdir(exist_ok=True)

    summary = {"files": {}, "errors": []}
    for src in sorted(src_root.iterdir()):
        if src.is_dir(): continue
        if src.suffix.lower() == ".md": continue  # skip existing markdown
        if src.name.startswith("."): continue
        ext = src.suffix.lower()
        conv = CONVERTERS.get(ext)
        if not conv:
            print(f"Skip (unsupported): {src.name}"); continue
        stem = safe_name(src.stem)
        img_dir = img_root / stem
        img_dir.mkdir(exist_ok=True)
        try:
            text, n_imgs = conv(src, img_dir)
            out_path = out_root / f"{stem}.md"
            out_path.write_text(text + "\n")
            # prune image dir if no images extracted
            if not any(img_dir.iterdir()):
                img_dir.rmdir()
            summary["files"][src.name] = {"output": str(out_path.name), "images": n_imgs}
            print(f"[ok ] {src.name} → {out_path.name} ({n_imgs} images)")
        except Exception as e:
            summary["errors"].append({"file": src.name, "error": str(e)})
            print(f"[err] {src.name}: {e}", file=sys.stderr)
    # Don't pollute the top-level with an intermediate summary file — the
    # final `_summary.md` / `_summary.json` written by cleanup.py supersedes
    # it. Keep the conversion-stage summary inside the hidden meta dir.
    meta_dir = out_root / ".pipeline-meta"; meta_dir.mkdir(exist_ok=True)
    (meta_dir / "conversion_summary.json").write_text(json.dumps(summary, indent=2))
    print(f"\nWrote {out_root}")


if __name__ == "__main__":
    main()
